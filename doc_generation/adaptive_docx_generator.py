# ============================================================
# YantraAI - Adaptive DOCX Generator
# JSON -> Microsoft Word (.docx)
# ============================================================

import json
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ============================================================
# HELPERS
# ============================================================

def safe_filename(name):
    name = str(name or "document")

    name = re.sub(
        r'[<>:"/\\|?*]',
        "_",
        name
    )

    return name.strip().replace(" ", "_")


def humanize_key(key):
    key = str(key)

    key = key.replace("_", " ")

    key = re.sub(
        r"([a-z])([A-Z])",
        r"\1 \2",
        key
    )

    return key.title()


def format_value(value):

    if value is None:
        return ""

    if isinstance(value, bool):
        return "Yes" if value else "No"

    if isinstance(value, (dict, list)):
        return json.dumps(
            value,
            ensure_ascii=False
        )

    return str(value)


# ============================================================
# ADAPTIVE DOCX GENERATOR
# ============================================================

class AdaptiveDocxGenerator:

    def __init__(self, data):

        self.data = data

        self.document = Document()

        formatting = (
            data.get("formatting", {})
        )

        # Also support formatting inside content
        if not formatting:

            content = data.get(
                "content",
                {}
            )

            if isinstance(content, dict):

                formatting = content.get(
                    "formatting",
                    {}
                )

        if not isinstance(
            formatting,
            dict
        ):
            formatting = {}

        self.font_name = formatting.get(
            "font",
            "Arial"
        )

        self.font_size = self.parse_font_size(
            formatting.get(
                "size",
                11
            )
        )

        self._setup_document()

    # ========================================================
    # DOCUMENT SETUP
    # ========================================================

    def _setup_document(self):

        style = self.document.styles["Normal"]

        style.font.name = (
            self.font_name
        )

        style.font.size = Pt(
            self.font_size
        )

    # ========================================================
    # FONT SIZE
    # ========================================================

    def parse_font_size(self, value):

        try:

            match = re.search(
                r"\d+(?:\.\d+)?",
                str(value)
            )

            if match:
                return float(
                    match.group()
                )

        except Exception:
            pass

        return 11

    # ========================================================
    # ADD TEXT
    # ========================================================

    def add_text(
        self,
        text,
        bold=False,
        italic=False,
        size=None,
        alignment=None
    ):

        paragraph = (
            self.document.add_paragraph()
        )

        if alignment:

            paragraph.alignment = (
                alignment
            )

        run = paragraph.add_run(
            format_value(text)
        )

        run.bold = bold
        run.italic = italic

        run.font.name = (
            self.font_name
        )

        run.font.size = Pt(
            size or self.font_size
        )

        return paragraph

    # ========================================================
    # HEADING
    # ========================================================

    def add_heading(
        self,
        text,
        level=1
    ):

        level = max(
            1,
            min(level, 9)
        )

        return self.document.add_heading(
            humanize_key(text),
            level=level
        )

    # ========================================================
    # TABLE
    # ========================================================

    def add_table(
        self,
        columns,
        rows,
        title=None
    ):

        if title:

            self.add_heading(
                title,
                2
            )

        if not columns:
            return

        table = (
            self.document.add_table(
                rows=1,
                cols=len(columns)
            )
        )

        table.style = "Table Grid"

        table.alignment = (
            WD_TABLE_ALIGNMENT.CENTER
        )

        # ----------------------------------------------------
        # HEADER
        # ----------------------------------------------------

        header_cells = (
            table.rows[0].cells
        )

        for index, column in enumerate(
            columns
        ):

            header_cells[index].text = (
                humanize_key(column)
            )

            for run in (
                header_cells[index]
                .paragraphs[0]
                .runs
            ):

                run.bold = True

        # ----------------------------------------------------
        # ROWS
        # ----------------------------------------------------

        for row in rows:

            cells = (
                table.add_row().cells
            )

            if isinstance(
                row,
                dict
            ):

                values = [
                    row.get(
                        column,
                        ""
                    )
                    for column in columns
                ]

            else:

                values = list(row)

            for index in range(
                len(columns)
            ):

                value = (
                    values[index]
                    if index < len(values)
                    else ""
                )

                cells[index].text = (
                    format_value(value)
                )

        self.document.add_paragraph()

    # ========================================================
    # DICTIONARY
    # ========================================================

    def render_dict(
        self,
        data,
        level=2
    ):

        for key, value in data.items():

            if key == "formatting":
                continue

            # Nested dictionary
            if isinstance(
                value,
                dict
            ):

                self.add_heading(
                    key,
                    level
                )

                self.render_dict(
                    value,
                    min(level + 1, 9)
                )

            # List
            elif isinstance(
                value,
                list
            ):

                self.render_list(
                    key,
                    value,
                    level
                )

            # Normal value
            else:

                paragraph = (
                    self.document.add_paragraph()
                )

                label = paragraph.add_run(
                    humanize_key(key)
                    + ": "
                )

                label.bold = True

                paragraph.add_run(
                    format_value(value)
                )

    # ========================================================
    # LIST
    # ========================================================

    def render_list(
        self,
        title,
        items,
        level=2
    ):

        if not items:
            return

        # ----------------------------------------------------
        # LIST OF DICTIONARIES
        # ----------------------------------------------------

        if all(
            isinstance(item, dict)
            for item in items
        ):

            flat = all(
                all(
                    not isinstance(
                        value,
                        (dict, list)
                    )
                    for value in item.values()
                )
                for item in items
            )

            if flat:

                columns = []

                for item in items:

                    for key in item:

                        if key not in columns:

                            columns.append(
                                key
                            )

                self.add_table(
                    columns,
                    items,
                    title=title
                )

                return

        # ----------------------------------------------------
        # SIMPLE LIST
        # ----------------------------------------------------

        self.add_heading(
            title,
            level
        )

        for item in items:

            if isinstance(
                item,
                (dict, list)
            ):

                self.render_dict(
                    item,
                    level + 1
                )

            else:

                paragraph = (
                    self.document.add_paragraph(
                        style="List Bullet"
                    )
                )

                paragraph.add_run(
                    format_value(item)
                )

    # ========================================================
    # SECTIONS
    # ========================================================

    def render_sections(
        self,
        sections
    ):

        for section in sections:

            if not isinstance(
                section,
                dict
            ):
                continue

            title = (
                section.get(
                    "section_title"
                )
                or section.get(
                    "title"
                )
                or section.get(
                    "name"
                )
            )

            if title:

                self.add_heading(
                    title,
                    1
                )

            content = section.get(
                "content"
            )

            if isinstance(
                content,
                dict
            ):

                self.render_dict(
                    content,
                    2
                )

            elif isinstance(
                content,
                list
            ):

                self.render_list(
                    "Content",
                    content,
                    2
                )

            elif content is not None:

                self.add_text(
                    content
                )

    # ========================================================
    # TABLES
    # ========================================================

    def render_tables(
        self,
        tables
    ):

        for table in tables:

            if not isinstance(
                table,
                dict
            ):
                continue

            columns = (
                table.get("columns")
                or table.get("headers")
                or []
            )

            rows = (
                table.get("rows")
                or table.get("data")
                or []
            )

            title = (
                table.get("table_title")
                or table.get("title")
                or table.get("name")
            )

            self.add_table(
                columns,
                rows,
                title
            )

    # ========================================================
    # MAIN RENDER
    # ========================================================

    def render(self):

        # ----------------------------------------------------
        # TITLE
        # ----------------------------------------------------

        metadata = self.data.get(
            "metadata",
            {}
        )

        if isinstance(
            metadata,
            dict
        ):

            title = metadata.get(
                "title"
            )

            if title:

                paragraph = (
                    self.document.add_paragraph()
                )

                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                )

                run = paragraph.add_run(
                    str(title)
                )

                run.bold = True

                run.font.name = (
                    self.font_name
                )

                run.font.size = Pt(
                    self.font_size + 5
                )

        # ----------------------------------------------------
        # CONTENT
        # ----------------------------------------------------

        content = self.data.get(
            "content",
            {}
        )

        if isinstance(
            content,
            dict
        ):

            # Sections
            if "sections" in content:

                self.render_sections(
                    content["sections"]
                )

            # Tables
            if "tables" in content:

                self.render_tables(
                    content["tables"]
                )

            # Everything else
            for key, value in content.items():

                if key in {
                    "sections",
                    "tables",
                    "formatting"
                }:
                    continue

                if isinstance(
                    value,
                    dict
                ):

                    self.add_heading(
                        key,
                        1
                    )

                    self.render_dict(
                        value,
                        2
                    )

                elif isinstance(
                    value,
                    list
                ):

                    self.render_list(
                        key,
                        value,
                        1
                    )

                else:

                    self.add_text(
                        f"{humanize_key(key)}: "
                        f"{format_value(value)}"
                    )

        elif isinstance(
            content,
            list
        ):

            self.render_list(
                "Content",
                content,
                1
            )

        else:

            self.add_text(
                content
            )

        return self.document

    # ========================================================
    # SAVE
    # ========================================================

    def save(
        self,
        output_path
    ):

        output_path = Path(
            output_path
        )

        output_path.parent.mkdir(
            parents=True,
            exist_ok=True
        )

        self.document.save(
            output_path
        )

        return output_path


# ============================================================
# PUBLIC FUNCTION
# ============================================================

def load_json(json_path):

    with open(
        json_path,
        "r",
        encoding="utf-8"
    ) as file:

        return json.load(file)


def generate_docx_from_json(
    json_path,
    output_path=None
):

    data = load_json(
        json_path
    )

    generator = AdaptiveDocxGenerator(
        data
    )

    generator.render()

    if output_path is None:

        base_dir = (
            Path(json_path).parent.parent
        )

        output_dir = (
            base_dir / "docx_output"
        )

        output_dir.mkdir(
            exist_ok=True
        )

        name = safe_filename(
            data.get(
                "document_name",
                Path(json_path).stem
            )
        )

        output_path = (
            output_dir
            / f"{name}.docx"
        )

    return generator.save(
        output_path
    )